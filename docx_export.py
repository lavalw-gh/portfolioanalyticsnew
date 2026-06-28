from __future__ import annotations

import sys
from pathlib import Path

import pandas as pd

try:
    from docx import Document
except ModuleNotFoundError:
    bundled_site_packages = Path(
        r"C:\Users\GW\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\Lib\site-packages"
    )
    if bundled_site_packages.exists():
        sys.path.insert(0, str(bundled_site_packages))
    from docx import Document

from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from .charts import make_chart_image


HOLDINGS_RIGHT_ALIGN_COLUMNS = {
    "Shares",
    "Cost (£)",
    "Value (£)",
    "Weight",
    "YTD Return",
    "Period Return",
}
KEY_METRICS_CENTER_COLUMNS = {"Portfolio Value", "Benchmark Value"}
MAX_PRICE_QUALITY_EXPORT_ROWS = 25
MAX_GENERIC_EXPORT_ROWS = 300
TABLE_WIDTH_CM = 19.0
HOLDINGS_TICKER_WIDTH_CM = 1.60
HOLDINGS_VALUE_WIDTH_CM = 2.00
PRICE_QUALITY_EXPLANATIONS = [
    {
        "Category": "Missing prices",
        "Action": "flagged",
        "Explanation": "Yahoo returned blank prices for one or more ticker/date points. These are reported so the data gap is visible.",
    },
    {
        "Category": "Forward fill",
        "Action": "carried forward",
        "Explanation": "A missing price was filled using the most recent previous valid price, within the allowed carry-forward limit.",
    },
    {
        "Category": "Unit regime",
        "Action": "corrected",
        "Explanation": "A likely 100x unit error was corrected, usually pence versus pounds for UK-listed securities.",
    },
    {
        "Category": "Residual spike",
        "Action": "corrected",
        "Explanation": "An isolated one-day price spike or drop was replaced using nearby prices when both sides supported the correction.",
    },
    {
        "Category": "Robust local",
        "Action": "flag_only",
        "Explanation": "A price looked unusual compared with nearby prices, but was not corrected automatically because confidence was lower.",
    },
]


def _is_holdings_value_column(column_name: str) -> bool:
    text = str(column_name).strip()
    return (
        text in HOLDINGS_RIGHT_ALIGN_COLUMNS
        or text.startswith("Cost (")
        or text.startswith("Value (")
    )


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, color: str = "999999", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))


def _set_table_width(table, width_cm: float = TABLE_WIDTH_CM) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_cm * 567)))


def _set_table_grid(table, widths_cm: list[float]) -> None:
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width_cm in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width_cm * 567)))
        tbl_grid.append(grid_col)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def _set_cell_text(cell, value, *, bold: bool = False, font_size: int = 8, align=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("" if value is None else str(value))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _set_table_autofit_off(table) -> None:
    table.autofit = False
    _set_table_width(table)
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def _column_widths(title: str, columns: list[str]) -> list[float]:
    if not columns:
        return [TABLE_WIDTH_CM]
    if title == "Holdings":
        widths: list[float] = []
        fixed_used = 0.0
        name_idx = None
        for i, col in enumerate(columns):
            if col == "Ticker":
                width = HOLDINGS_TICKER_WIDTH_CM
            elif col == "Name":
                width = 0.0
                name_idx = i
            elif _is_holdings_value_column(col):
                width = HOLDINGS_VALUE_WIDTH_CM
            else:
                width = 1.6
            widths.append(width)
            fixed_used += width
        if name_idx is not None:
            widths[name_idx] = max(3.0, TABLE_WIDTH_CM - fixed_used)
        return widths
    if title == "Price data quality":
        preferred = {
            "Category": 2.2,
            "Symbol": 1.5,
            "Date / Range": 2.4,
            "Issue": 4.2,
            "Action": 2.0,
            "Old Price": 1.5,
            "New Price": 1.5,
            "Method": 2.4,
            "Confidence": 1.3,
        }
        widths = [preferred.get(str(col), 2.0) for col in columns]
        total = sum(widths)
        if total > TABLE_WIDTH_CM:
            scale = TABLE_WIDTH_CM / total
            widths = [max(1.1, w * scale) for w in widths]
        return widths
    return [TABLE_WIDTH_CM / len(columns) for _ in columns]


def _style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in (("Heading 1", 14), ("Heading 2", 11)):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(15, 23, 42)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)


def _add_paragraph(doc: Document, text: str = ""):
    paragraph = doc.add_paragraph("" if text is None else str(text))
    paragraph.style = doc.styles["Normal"]
    return paragraph


def _add_table(doc: Document, title: str, df: pd.DataFrame) -> None:
    doc.add_heading(title, level=2)
    if df is None:
        raw_df = pd.DataFrame()
    else:
        raw_df = pd.DataFrame(df).copy()

    if title == "Price data quality" and not raw_df.empty:
        if "Action" in raw_df.columns:
            raw_df = raw_df[~raw_df["Action"].astype(str).str.lower().isin({"no_change", "checked", "none"})]
        if len(raw_df) > MAX_PRICE_QUALITY_EXPORT_ROWS:
            group_cols = [col for col in ("Category", "Action") if col in raw_df.columns]
            if group_cols:
                raw_df = (
                    raw_df.groupby(group_cols, dropna=False)
                    .size()
                    .reset_index(name="Count")
                    .sort_values("Count", ascending=False)
                )
            else:
                raw_df = pd.DataFrame({
                    "Category": ["Price data"],
                    "Issue": [f"{len(raw_df)} price-quality rows detected"],
                })
            if len(raw_df) > MAX_PRICE_QUALITY_EXPORT_ROWS:
                raw_df = raw_df.head(MAX_PRICE_QUALITY_EXPORT_ROWS)

    df2 = raw_df.fillna("").astype(str)
    original_rows = len(df2)
    if title == "Price data quality" and original_rows > MAX_PRICE_QUALITY_EXPORT_ROWS:
        df2 = df2.head(MAX_PRICE_QUALITY_EXPORT_ROWS).copy()
        summary = {col: "" for col in df2.columns}
        if "Category" in summary:
            summary["Category"] = "Summary"
        if "Issue" in summary:
            summary["Issue"] = (
                f"{original_rows - MAX_PRICE_QUALITY_EXPORT_ROWS} additional price-quality rows omitted from Word export."
            )
        df2 = pd.concat([df2, pd.DataFrame([summary])], ignore_index=True)
    elif original_rows > MAX_GENERIC_EXPORT_ROWS:
        df2 = df2.head(MAX_GENERIC_EXPORT_ROWS).copy()
        summary = {col: "" for col in df2.columns}
        first_col = df2.columns[0] if len(df2.columns) else None
        if first_col is not None:
            summary[first_col] = f"{original_rows - MAX_GENERIC_EXPORT_ROWS} additional rows omitted from Word export."
        df2 = pd.concat([df2, pd.DataFrame([summary])], ignore_index=True)
    columns = list(df2.columns) if not df2.empty else []
    widths = _column_widths(title, columns)
    row_count = max(1, len(df2)) + 1
    col_count = max(1, len(columns))

    table = doc.add_table(rows=row_count, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_autofit_off(table)
    _set_table_grid(table, widths)
    _set_repeat_table_header(table.rows[0])

    if columns:
        for col_idx, col in enumerate(columns):
            cell = table.cell(0, col_idx)
            align = None
            if title == "Key metrics" and col in KEY_METRICS_CENTER_COLUMNS:
                align = WD_ALIGN_PARAGRAPH.CENTER
            elif title == "Holdings" and _is_holdings_value_column(col):
                align = WD_ALIGN_PARAGRAPH.RIGHT
            _set_cell_width(cell, widths[col_idx])
            _set_cell_shading(cell, "DBEAFE")
            _set_cell_border(cell, color="666666")
            _set_cell_text(cell, col, bold=True, font_size=8, align=align)
    else:
        cell = table.cell(0, 0)
        _set_cell_width(cell, TABLE_WIDTH_CM)
        _set_cell_shading(cell, "DBEAFE")
        _set_cell_border(cell, color="666666")
        _set_cell_text(cell, "No data", bold=True, font_size=8)

    if df2.empty or not columns:
        cell = table.cell(1, 0)
        _set_cell_width(cell, TABLE_WIDTH_CM)
        _set_cell_border(cell)
        _set_cell_text(cell, "No data", font_size=8)
    else:
        for row_idx, (_, row_vals) in enumerate(df2.iterrows(), start=1):
            is_total_row = title == "Holdings" and str(row_vals.get("Ticker", "")).strip() == "TOTAL"
            for col_idx, col in enumerate(columns):
                cell = table.cell(row_idx, col_idx)
                align = None
                if title == "Key metrics" and col in KEY_METRICS_CENTER_COLUMNS:
                    align = WD_ALIGN_PARAGRAPH.CENTER
                elif title == "Holdings" and _is_holdings_value_column(col):
                    align = WD_ALIGN_PARAGRAPH.RIGHT
                _set_cell_width(cell, widths[col_idx])
                _set_cell_border(cell)
                _set_cell_text(
                    cell,
                    row_vals.get(col, ""),
                    bold=is_total_row,
                    font_size=8 if title in ("Holdings", "Price data quality", "Price data quality explanations") else 10,
                    align=align,
                )

    _add_paragraph(doc, "")


def _add_price_quality_explanation(doc: Document) -> None:
    _add_table(doc, "Price data quality explanations", pd.DataFrame(PRICE_QUALITY_EXPLANATIONS))


def _add_page_break_before_next_heading(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)


def write_docx(results: dict, output_path: str):
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    tmp_dir = output_file.parent / f".{output_file.stem}_assets"
    tmp_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _style_document(doc)

    doc.add_heading("Portfolio Analysis", level=1)
    _add_paragraph(doc, f"Benchmark: {results.get('benchmark', '')}")
    _add_paragraph(doc, f"Date range: {results.get('start_effective')} to {results.get('end_effective')}")
    _add_paragraph(doc, "")
    _add_table(doc, "Price data quality", results.get("price_quality_df", pd.DataFrame()))
    _add_price_quality_explanation(doc)

    bench_cum = results.get("bench_cum_window", pd.Series(dtype=float))
    for i, pname in enumerate(results.get("selected_names", [])):
        if i > 0:
            _add_page_break_before_next_heading(doc)
        pdata = results.get("portfolio_data", {}).get(pname, {})
        doc.add_heading(f"Portfolio: {pname}", level=1)

        chart_path = make_chart_image(pname, pdata, tmp_dir, bench_cum_window=bench_cum)
        if chart_path.exists():
            doc.add_picture(str(chart_path), width=Cm(TABLE_WIDTH_CM))
            _add_paragraph(doc, "")

        _add_table(doc, "Key metrics", pdata.get("metrics_df", pd.DataFrame()))
        _add_table(doc, "Holdings", pdata.get("holdings_df", pd.DataFrame()))

    doc.save(str(output_file))
